#!/usr/bin/env python3
"""Extract searchable text from an engineering project folder.

This helper is intentionally conservative:
- It never edits source files.
- Converted files and text outputs are written under the chosen output folder.
- Unsupported files are skipped with an error entry in the manifest.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
from pathlib import Path


SUPPORTED = {".docx", ".doc", ".docm", ".wps", ".xlsx", ".xls", ".pdf", ".txt"}


def safe_name(path: Path, root: Path) -> str:
    rel = str(path.relative_to(root))
    return re.sub(r"[^\w.\-\u4e00-\u9fff]+", "_", rel)


def read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = " ".join(para.text.split())
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            vals = [" ".join(cell.text.split()) for cell in row.cells]
            if any(vals):
                lines.append("\t".join(vals))
    return "\n".join(lines)


def read_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=True, read_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        lines.append(f"### SHEET {ws.title}")
        for row in ws.iter_rows(values_only=True):
            vals = [str(v).strip() if v is not None else "" for v in row]
            if any(vals):
                lines.append("\t".join(vals))
    return "\n".join(lines)


def libreoffice_convert(src: Path, out_dir: Path, fmt: str) -> Path | None:
    out_dir.mkdir(parents=True, exist_ok=True)
    cmd = ["libreoffice", "--headless", "--convert-to", fmt, "--outdir", str(out_dir), str(src)]
    subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=120)
    candidates = sorted(out_dir.glob(f"{src.stem}.*"))
    return candidates[0] if candidates else None


def pdftotext(src: Path, out_file: Path) -> bool:
    exe = shutil.which("pdftotext")
    if not exe:
        return False
    subprocess.run([exe, "-layout", str(src), str(out_file)], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=180)
    return out_file.exists()


def extract_file(path: Path, root: Path, text_dir: Path, converted_dir: Path) -> dict:
    ext = path.suffix.lower()
    record = {"source": str(path), "status": "ok", "text_file": None, "error": None}
    out_file = text_dir / f"{safe_name(path, root)}.txt"
    try:
        if ext == ".txt":
            text = path.read_text(encoding="utf-8", errors="ignore")
        elif ext == ".docx":
            text = read_docx(path)
        elif ext in {".doc", ".docm", ".wps"}:
            converted = libreoffice_convert(path, converted_dir / safe_name(path, root), "docx")
            if not converted:
                converted = libreoffice_convert(path, converted_dir / safe_name(path, root), "txt:Text")
            if not converted:
                raise RuntimeError("LibreOffice conversion produced no output")
            text = read_docx(converted) if converted.suffix.lower() == ".docx" else converted.read_text(encoding="utf-8", errors="ignore")
        elif ext == ".xlsx":
            text = read_xlsx(path)
        elif ext == ".xls":
            converted = libreoffice_convert(path, converted_dir / safe_name(path, root), "xlsx")
            if not converted:
                raise RuntimeError("LibreOffice .xls conversion produced no output")
            text = read_xlsx(converted)
        elif ext == ".pdf":
            if not pdftotext(path, out_file):
                raise RuntimeError("pdftotext unavailable or failed")
            record["text_file"] = str(out_file)
            return record
        else:
            record["status"] = "skipped"
            record["error"] = f"unsupported extension: {ext}"
            return record
        out_file.write_text(text, encoding="utf-8", errors="ignore")
        record["text_file"] = str(out_file)
    except Exception as exc:  # noqa: BLE001
        record["status"] = "error"
        record["error"] = repr(exc)
    return record


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract searchable text from project documents.")
    parser.add_argument("project_folder", type=Path)
    parser.add_argument("--out", type=Path, required=True, help="Output folder for text corpus and manifest.")
    args = parser.parse_args()

    root = args.project_folder.resolve()
    text_dir = args.out.resolve() / "text"
    converted_dir = args.out.resolve() / "converted"
    text_dir.mkdir(parents=True, exist_ok=True)
    converted_dir.mkdir(parents=True, exist_ok=True)

    files = [p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED]
    manifest = [extract_file(p, root, text_dir, converted_dir) for p in files]
    (args.out / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"files": len(files), "ok": sum(1 for x in manifest if x["status"] == "ok"), "out": str(args.out)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
