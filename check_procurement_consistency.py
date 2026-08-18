#!/usr/bin/env python3
"""Check consistency between procurement Excel files and product-description DOCX files."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

CN_NUM = "一二三四五六七八九十"

# 关键类边界护栏定义（类名 → 允许辅助词，用于「类名自洽」抽查）
# 注意：这些类名来自横沟/选煤厂案例，仅当当前工程真实生成了同名分类时才生效（护栏）；
# 绝不反向推导——不要求工程必须存在这些分类。分类由当前工程物料决定（见 SKILL.md 动态分类方法论）。
CATEGORY_ALLOWED = {
    "工控服务器 UPS": ["服务器", "工控机", "工作站", "笔记本", "Pad", "硬盘", "超融合", "UPS", "授时"],
    "网络交换设备": ["交换机", "网关", "防火墙"],
    "光纤网线类": ["光纤", "光缆", "网线", "跳线", "收发器"],
    "机柜箱体配电": ["机柜", "柜", "箱", "机架", "模块"],
}

# 安装工序行特征（名称以这些词结尾视为疑似工序行，需人工复核，非设备材料）
# 注意：不能机械删除所有含"安装"的名称（如"安装辅材""安装套件"是材料）；
# 本脚本只标记疑似项，最终判断需结合单位、数量、参数、上下文（见 SKILL.md 原始资料提取规则）。
INSTALL_SUFFIX = ("安装", "敷设")


def norm(value) -> str:
    return str(value).strip() if value is not None else ""


def read_items(xlsx: Path) -> list[dict]:
    from openpyxl import load_workbook

    wb = load_workbook(xlsx, data_only=True, read_only=True)
    ws = wb.active
    rows: list[dict] = []
    last_name = ""
    for row in ws.iter_rows(values_only=True):
        vals = [norm(v) for v in row[:5]]
        if len(vals) < 5 or not any(vals):
            continue
        if vals[1] in {"设备名称", "设备材料名称"}:
            continue
        if vals[0] and not any(vals[1:]):
            continue
        if vals[1]:
            last_name = vals[1]
            rows.append({"name": vals[1], "spec": vals[2], "qty": vals[3], "unit": vals[4]})
        elif last_name and (vals[2] or vals[3] or vals[4]):
            rows.append({"name": last_name, "spec": vals[2], "qty": vals[3], "unit": vals[4]})
    return rows


def read_doc(docx: Path) -> tuple[list[str], list[str]]:
    from docx import Document

    doc = Document(str(docx))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading") and p.text.strip()]
    return paras, headings


def section_numbers_ok(headings: list[str]) -> tuple[bool, list[str], list[str]]:
    got: list[str] = []
    for heading in headings[1:]:
        match = re.match(r"([一二三四五六七八九十]+)、", heading)
        if match:
            got.append(match.group(1))
    expected = list(CN_NUM[: len(got)])
    return got == expected, got, expected


def render_docx(docx: Path, out_dir: Path, renderer: str | None) -> tuple[int | None, int]:
    if not renderer:
        return None, 0
    out_dir.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        [sys.executable, renderer, str(docx), "--output_dir", str(out_dir)],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        timeout=120,
    )
    return result.returncode, len(list(out_dir.glob("*.png")))


def main() -> int:
    parser = argparse.ArgumentParser(description="Check procurement Excel and product-description Word consistency.")
    parser.add_argument("procurement_folder", type=Path)
    parser.add_argument("--source-text", type=Path, default=None, help="Optional project source text to check referenced terms.")
    parser.add_argument("--renderer", default=None, help="Optional render_docx.py path for DOCX visual render smoke check.")
    parser.add_argument("--source-count", type=int, default=None, help="Optional source device-row count for conservation check (classified rows must not exceed it).")
    parser.add_argument("--out", type=Path, required=True, help="Output JSON report path.")
    args = parser.parse_args()

    root = args.procurement_folder.resolve()
    source_text = args.source_text.read_text(encoding="utf-8", errors="ignore") if args.source_text and args.source_text.exists() else ""
    compact_source = re.sub(r"\s+", "", source_text)
    render_root = args.out.parent / "consistency_render"
    if render_root.exists():
        shutil.rmtree(render_root)

    reports: list[dict] = []
    for folder in sorted([p for p in root.iterdir() if p.is_dir()]):
        xlsx_files = sorted(folder.glob("*.xlsx"))
        docx_files = sorted(folder.glob("*产品说明.docx"))
        if not xlsx_files or not docx_files:
            continue
        xlsx = xlsx_files[0]
        docx = docx_files[0]
        items = read_items(xlsx)
        install_rows = [item["name"] for item in items if item["name"].endswith(INSTALL_SUFFIX)]
        boundary_issues: list[str] = []
        allowed = CATEGORY_ALLOWED.get(folder.name)
        if allowed:
            for item in items:
                if not any(kw in item["name"] for kw in allowed):
                    boundary_issues.append(item["name"])
        paras, headings = read_doc(docx)
        text = "\n".join(paras)

        declared_match = re.search(r"本分类清单共包含\s*(\d+)\s*项采购条目", text)
        declared_count = int(declared_match.group(1)) if declared_match else None

        listed_match = re.search(r"主要设备和材料包括：(.+?)。具体数量", text)
        unknown_names: list[str] = []
        if listed_match:
            listed = [name.strip().removesuffix("等") for name in listed_match.group(1).split("、") if name.strip()]
            excel_names = {item["name"] for item in items}
            unknown_names = [name for name in listed if name not in excel_names]

        nums_ok, got_nums, expected_nums = section_numbers_ok(headings)
        render_code, render_pages = render_docx(docx, render_root / docx.stem, args.renderer)

        source_terms = []
        unsupported_terms = []
        if source_text:
            for term in ["KYN28", "IEC61850", "OPGW", "2x2M", "TCP/IP", "火灾自动报警系统", "防雷", "接地", "UPS", "PLC", "工业以太网", "调度数据网"]:
                if term in text:
                    if re.sub(r"\s+", "", term) in compact_source:
                        source_terms.append(term)
                    else:
                        unsupported_terms.append(term)

        reports.append(
            {
                "category": folder.name,
                "xlsx_rows": len(items),
                "declared_count": declared_count,
                "count_ok": declared_count == len(items),
                "unknown_listed_names": unknown_names,
                "section_numbers_ok": nums_ok,
                "section_numbers": got_nums,
                "expected_section_numbers": expected_nums,
                "bad_punctuation": [p for p in ["、，", "，，"] if p in text],
                "source_supported_terms": source_terms,
                "source_unsupported_terms": unsupported_terms,
                "install_rows": install_rows,
                "boundary_issues": boundary_issues,
                "render_returncode": render_code,
                "render_pages": render_pages,
            }
        )

    total_classified = sum(r["xlsx_rows"] for r in reports)
    conservation = None
    if args.source_count is not None:
        conservation = {
            "source_count": args.source_count,
            "classified_rows": total_classified,
            "ok": total_classified <= args.source_count,
        }

    summary = {
        "total_docs": len(reports),
        "count_mismatches": [r["category"] for r in reports if not r["count_ok"]],
        "unknown_listed_name_docs": [r["category"] for r in reports if r["unknown_listed_names"]],
        "section_number_errors": [r["category"] for r in reports if not r["section_numbers_ok"]],
        "bad_punctuation_docs": [r["category"] for r in reports if r["bad_punctuation"]],
        "source_unsupported_docs": [r["category"] for r in reports if r["source_unsupported_terms"]],
        "render_errors": [r["category"] for r in reports if r["render_returncode"] not in (None, 0) or (r["render_returncode"] == 0 and r["render_pages"] == 0)],
        "install_row_categories": [r["category"] for r in reports if r["install_rows"]],
        "boundary_issue_categories": [r["category"] for r in reports if r["boundary_issues"]],
        "conservation": conservation,
    }
    result = {"summary": summary, "reports": reports}
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
